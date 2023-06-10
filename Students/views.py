from django.shortcuts import render, redirect
from django.http import HttpResponse
from Students.models import Student
from .forms import StudentForm
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.contrib.auth.models import User
import random
from Classes.models import Class
import string
from Enrollments.models import Enrollment
from django.contrib.auth.hashers import make_password
import docx
from django.http import FileResponse
import io
from docx.enum.text import WD_ALIGN_PARAGRAPH
from FeesManagement.templatetags.fees_tags import fees_balance
from django.core.paginator import Paginator
from Classes.forms import ClassForm
from django import forms
import plotly.graph_objs as go
import plotly
import qrcode
import os
import requests
import json
from PIL import Image, ImageDraw, ImageFont


# cPanel API credentials
api_username = "edumetrics"
api_token = "B4Z1XQX2I50FLRB57QE6XZA8WEN0K44I"
domain = "edu-metrics.com"
api_url = f"https://edu-metrics.com:2083/execute/Email/add_pop"


def generate_student_id():
    # Generate a 4-letter random string
    letters = string.ascii_uppercase
    random_letters = ''.join(random.choice(letters) for i in range(4))
    
    # Generate a 4-digit random number
    random_number = str(random.randint(1000, 9999))
    
    # Combine the two to form a student ID
    student_id = random_letters + random_number
    
    return student_id


def get_qr(request, student):
    student = Student.objects.get(id=student)
    qr = qrcode.QRCode(
        version=1,
        box_size=10,
        border=5
    )
    data = f'''
    First Name: {student.first_name},
    Last Name: {student.last_name},
    Student Id: {student.student_id},
    School: {student.school},
    Class: {student.Class}
    '''
    qr.add_data(data)
    qr.make(fit=True)
    qr_img = qr.make_image(fill_color='White', back_color='#103741')
    img = qr_img.get_image()
    img = img.convert('RGB')

    response = HttpResponse(content_type='image/png')
    img.save(response, 'PNG')
    return response


def download_students_info(request):
    doc = docx.Document()
    doc.add_heading(f'{request.user.schooladministrator.current_school} Students Info', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    table = doc.add_table(rows=request.user.schooladministrator.current_school.students.count()+1, cols=6)
    table.cell(0, 0).text = 'First Name'
    table.cell(0, 1).text = 'Last_Name'
    table.cell(0, 2).text = 'Class'
    table.cell(0, 3).text = 'Student Id'
    table.cell(0, 4).text = 'Date Enrolled'
    table.cell(0, 5).text = 'Fees Balance'
    row = 1
    for student in request.user.schooladministrator.current_school.students.all():
        table.cell(row, 0).text = student.first_name
        table.cell(row, 1).text = student.last_name
        table.cell(row, 2).text = student.Class.Name
        table.cell(row, 3).text = student.student_id
        table.cell(row, 4).text = str(student.active_enrollment.Date)
        table.cell(row, 5).text = f'{fees_balance(student)} UGX'
        row += 1

    document_io = io.BytesIO()
    doc.save(document_io)
    document_io.seek(0)
    response = FileResponse(document_io, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{request.user.schooladministrator.current_school}_students_info.docx"'
    return response


def student_dashboard(request):
    return render(request, 'students_dashboard.html')


@login_required
def select_student_class(request):
    form = StudentForm()
    class_field = forms.ModelChoiceField(queryset=request.user.schooladministrator.current_school.classes.all(), widget=forms.Select(attrs={'class': 'form-control'}))
    form.fields = {
        'Class': class_field
    }
    context = {
        'form': form
    }
    if request.method == 'POST':
        form = StudentForm(request.POST)
        form.fields = {
            'Class': class_field
        }
        if form.is_valid():
            clas = form.cleaned_data['Class']
            return redirect('enroll-student', clas=clas.id)
    return render(request, 'select_student_class.html', context)


@login_required
def enroll_student(request, clas):
    clas = request.user.schooladministrator.current_school.classes.get(id=clas)
    form = StudentForm(initial={'student_id': generate_student_id()})
    form.fields['Subjects'].queryset = clas.Level.Subjects.all()
    if clas.Streams.all():
        form.fields['Stream'].queryset = clas.Streams.all()
    else:
        del form.fields['Stream']
    if clas.Level.Name != 'Advance Level':
        del form.fields['Combination']
    if request.method == 'POST':
        form = StudentForm(request.POST, request.FILES)
        if clas.Level.Name != 'Advance Level':
            del form.fields['Combination']
        if not clas.Streams.all():
            del form.fields['Stream']
        if form.is_valid():
            student = form.save(commit=False)
            student.school = request.user.schooladministrator.current_school
            student.user = User.objects.create(username=form.cleaned_data['student_id'])
            student.Level = clas.Level
            student.save()
            request.user.schooladministrator.current_school.students.add(student)
            clas.Students.add(student)
            subjects = form.cleaned_data['Subjects']
            student.Subjects.set(subjects)
            student.Class = clas
            student.save()
            student.Level.Students.add(student)
            for subject in student.Subjects.all():
                subject.Students.add(student)
            enrollment = Enrollment.objects.create(School=request.user.schooladministrator.current_school, Student=student, Programme=student.Programme, By=request.user.schooladministrator)
            student.active_enrollment = enrollment
            student.save()
            student.Programme.Students.add(student)
            student.school.Enrollments.add(enrollment)
            messages.success(request, f'Student {student.first_name} {student.last_name} has been enrolled successfully.')
            return redirect('students')
    return render(request, 'enroll_student.html', {'form': form})


def change_student(request, student):
    student = request.user.schooladministrator.current_school.students.get(pk=student)
    form = StudentForm(instance=student)
    form.fields['Subjects'].queryset = student.Class.Level.Subjects.all()
    if student.Class.Streams.all():
        form.fields['Stream'].queryset = student.Class.Streams.all()
    else:
        del form.fields['Stream']
    if student.Class.Level.Name != 'Advance Level':
        del form.fields['Combination']
    context = {
        'form': form
    }
    if request.method == 'POST':
        form = StudentForm(request.POST, request.FILES, instance=student)
        if student.Class.Streams.all():
            form.fields['Stream'].queryset = student.Class.Streams.all()
        else:
            del form.fields['Stream']
        if student.Class.Level.Name != 'Advance Level':
            del form.fields['Combination']
        if form.is_valid():
            student = form.save()
            messages.success(request, f'Student {student.first_name} {student.last_name} info was changed successfully.')
            return redirect('students')
        else:
            print(form)
    else:
        return render(request, 'change_student.html', context)


def  import_student(request):
    return render(request, 'import_student.html')


def child_info(request, child):
    child = Student.objects.get(pk=child)
    context = {
        'child': child
    }
    return render(request, 'child_info.html', context)


def edit_student(request, student_id):
    student = get_object_or_404(Student, id=student_id)
    if request.method == 'POST':
        form = StudentForm(request.POST, instance=student)
        if form.is_valid():
            student = form.save(commit=False)
            student.school = request.user.teacher.school
            student.save()
            messages.success(request, 'Student updated successfully')
            return redirect('student_list')
    else:
        form = StudentForm(instance=student)
    return render(request, 'edit_student.html', {'form': form, 'student': student})

def delete_student(request, student_id):
    student = get_object_or_404(Student, id=student_id)
    if request.method == 'POST':
        student.delete()
        messages.success(request, 'Student deleted successfully')
        return redirect('student_list')
    return render(request, 'delete_student.html', {'student': student})

def student_list(request):
    students = request.user.schooladministrator.current_school.students.all().order_by('active_enrollment__Date').reverse()
    paginator = Paginator(students, 6)
    page = request.GET.get('page')
    students = paginator.get_page(page)
    context = {
        'students': students
    }
    return render(request, 'student_list.html', context)


def student_profile(request, student):
    student = request.user.schooladministrator.current_school.students.get(pk=student)
    x = [subject.name for subject in student.Subjects.all()]
    y1 = [70, 50, 45, 20, 34, 56]
    trace1 = go.Bar(x=x, y=y1, name='Average Score', marker=dict(color='#103741'))
    layout = go.Layout(title=f'Performance Statistics', xaxis=dict(title='Subjects'), yaxis=dict(title='Average Statistics'))
    chart = plotly.offline.plot({
        "data": [trace1,],
        "layout": layout
    }, output_type="div")
    context = {
        'student': student,
        'stats': chart
    }
    return render(request, 'student_profile.html', context)


def old_students(request):
    return render(request, 'old_students.html')

def terminate_student(request, student):
    student = request.user.schooladministrator.current_school.students.get(id=student)
    context = {
        'student': student
    }
    if request.method == 'POST':
        reason = request.POST['reason']
        request.user.schooladministrator.current_school.students.remove(student)
        student.education_history.create(
            Student=student,
            School=request.user.schooladministrator.current_school,
            From=student.active_enrollment.Date,
        )
        student.active_enrollment.Programme.Students.remove(student)
        student.Class.Students.remove(student)
        student.active_enrollment = None
        student.Class = None
        student.school = None
        student.save()
        return redirect('students')
    else:
        return render(request, 'terminate_student.html', context)
    

def search_student(request):
    form = StudentForm()
    fields = {
        'Names': form.fields['first_name'],
        'Gender': form.fields['Gender']
    }
    form.fields = fields
    context = {
        'form': form
    }
    return render(request, 'search_student.html', context)


def setup_account(request, student):
    student = request.user.schooladministrator.current_school.students.get(id=student)
    context = {
        'student': student
    }
    if request.method == 'POST':
        email = request.POST['email']
        email_password = request.POST['password']
        payload = {
            "email": email,
            "password": email_password,
            "quota": 250 # Mailbox quota in megabytes
        }

        # Request headers
        headers = {
            "Authorization": f"cpanel {api_username}:{api_token}",
            "Content-Type": "application/json"
        }

        # Send API request
        response = requests.post(api_url, data=json.dumps(payload), headers=headers)

        # Check the response
        if response.status_code == 200:
            response = json.dumps(response.text)
            student.user.email = email
            student.user.username = email
            student.user.password = make_password(email_password)
            student.user.save()
            messages.success(request, 'Student Account set up successfully.')
            return redirect('student-profile', student=student.id)
        else:
            return HttpResponse(f"An error occurred: {response.text}")
    return render(request, 'set_up_student_account.html', context)


def generate_idcard(request, student):
    student = request.user.schooladministrator.current_school.students.get(id=student)
    # Define the size of the ID card (in pixels)
    card_width = 800
    card_height = 500

    # Create a blank ID card with white background
    id_card = Image.new("RGB", (card_width, card_height), (255, 255, 255))
    draw = ImageDraw.Draw(id_card)

    # Load fonts
    font_header = ImageFont.truetype("./ARIAL.TTF", 40)
    font_motto = ImageFont.truetype("./ARIAL.TTF", 30)
    font_text = ImageFont.truetype("./ARIAL.TTF", 30)
    font_generated_by = ImageFont.truetype("./ARIAL.TTF", 20)

    # Add badge to the upper left corner with white background
    badge = Image.open("../badge.png")
    badge = badge.resize((100, 100))
    badge_with_bg = Image.new("RGB", badge.size, (255, 255, 255))
    badge_with_bg.paste(badge, (0, 0), mask=badge)
    id_card.paste(badge_with_bg, (50, 50))



    # Add "Mothercare Preparatory School" text
    school_name = student.school.name
    school_name_width, school_name_height = draw.textsize(school_name, font=font_header)
    school_name_x = 200
    school_name_y = 70
    draw.text((school_name_x, school_name_y), school_name, fill=(0, 0, 0), font=font_header)

    # Add line between the school name and motto
    line_start = (school_name_x, school_name_y + school_name_height + 5)
    line_end = (school_name_x + school_name_width, school_name_y + school_name_height + 5)
    draw.line([line_start, line_end], fill=(0, 0, 0), width=2)

    motto = "Inspiring Young Minds"
    motto_width, motto_height = draw.textsize(motto, font=font_motto)
    motto_x = school_name_x + (school_name_width - motto_width) // 2
    motto_y = school_name_y + school_name_height + 10
    draw.text((motto_x, motto_y), motto, fill=(0, 0, 0), font=font_motto)

    # Add student photo to ID card (left corner)
    student_photo = Image.open(student.photo.path)
    student_photo = student_photo.resize((200, 200))
    id_card.paste(student_photo, (50, 180))

    # Create a border around the student photo
    border_color = (254, 93, 55)
    border_width = 3
    bordered_photo = Image.new("RGB", (student_photo.width + 2 * border_width, student_photo.height + 2 * border_width), border_color)
    bordered_photo.paste(student_photo, (border_width, border_width))

    # Paste the bordered student photo onto the ID card
    id_card.paste(bordered_photo, (50, 180))

    # Calculate the proportional height of the text based on the photo height
    text_height = int(student_photo.height * 0.8)

    # Define text positions and offsets
    text_x = 300
    text_y = 180 + (student_photo.height - text_height) // 2
    line_height = int(text_height / 5)

    # Add student details to ID card
    draw.text((text_x, text_y), "Names: " + f'{student.first_name} {student.last_name}', fill=(0, 0, 0), font=font_text)
    draw.text((text_x, text_y + line_height), "ID: " + student.student_id, fill=(0, 0, 0), font=font_text)
    draw.text((text_x, text_y + 2 * line_height), "Class: " + student.Class.Name, fill=(0, 0, 0), font=font_text)

    # Add date issued and expiry date
    date_issued = "Date Issued: 2023-06-10"
    expiry_date = "Expiry Date: 2024-06-10"
    draw.text((text_x, text_y + 3 * line_height), date_issued, fill=(0, 0, 0), font=font_text)
    draw.text((text_x, text_y + 4 * line_height), expiry_date, fill=(0, 0, 0), font=font_text)

    border_width = 6
    draw.rectangle([(0, 0), (card_width - 1, card_height - 1)], outline=(254, 93, 55), width=border_width)


    # Add "Generated by Edumetrics" line
    generated_by = "Generated by Edumetrics"
    generated_by_width, generated_by_height = draw.textsize(generated_by, font=font_generated_by)
    generated_by_x = (card_width - generated_by_width) // 2
    generated_by_y = card_height - 50
    draw.text((generated_by_x, generated_by_y), generated_by, fill=(0, 0, 0), font=font_generated_by)

    # Save the ID card
    id_c = id_card.save("student_id_card.png")

    return FileResponse(open("student_id_card.png", 'rb'))