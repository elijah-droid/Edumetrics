from django.shortcuts import render, redirect, get_object_or_404
from django.http import FileResponse
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.utils import timezone
from .models import Examination
from .forms import ExamScheduleForm, PaperForm
import io
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.core.paginator import Paginator


@login_required
def schedule_exam(request):
    if request.method == 'POST':
        form = ExamScheduleForm(request.POST)
        if form.is_valid():
            exam_set = form.save(commit=False)
            exam_set.School = request.user.schooladministrator.current_school
            exam_set.save()
            request.user.schooladministrator.current_school.Examinations.add(exam_set)
            messages.success(request, 'Examination set scheduled successfully!')
            return redirect('examinations')
    else:
        form = ExamScheduleForm()
        form.fields['Classes'].queryset = request.user.schooladministrator.current_school.classes.all()
    context = {'form': form}
    return render(request, 'schedule_exam.html', context)


@login_required
def edit_examination_set(request, pk):
    examination_set = get_object_or_404(ExaminationSet, pk=pk)
    form = ExaminationSetForm(request.POST or None, instance=examination_set)
    if request.method == 'POST':
        if form.is_valid():
            form.save()
            return redirect('examination_set_detail', pk=pk)
    context = {
        'form': form,
        'examination_set': examination_set,
    }
    return render(request, 'edit_examination_set.html', context)

@login_required
def delete_examination(request, exam):
    examination_set = get_object_or_404(Examination, pk=exam)
    examination_set.delete()
    return redirect('examinations')


def examination_list(request):
    examinations = Examination.objects.all().order_by('Date')
    paginator = Paginator(examinations, 12)
    page = request.GET.get('page')
    examinations = paginator.get_page(page)
    return render(request, 'examinations.html', {'examinations': examinations})


def view_exam_reports(request, exam):
    exam = Examination.objects.get(pk=exam)
    content = {
        'reports': exam.Reports.all()
    }
    return render(request, 'reports.html', content)


def assessments(request, exam):
    exam = request.user.schooladministrator.current_school.Examinations.get(pk=exam)
    doc = docx.Document()
    for clas in request.user.schooladministrator.current_school.classes.all():
        doc.add_heading(f'{request.user.schooladministrator.current_school} {clas} {exam} Assessment Sheet', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        table = doc.add_table(rows=request.user.schooladministrator.current_school.students.count()+1, cols=request.user.schooladministrator.current_school.Subjects.count()+1)
        table.cell(0, 0).text = 'Student'
        cl = 1
        for subject in request.user.schooladministrator.current_school.Subjects.all():
            table.cell(0, cl).text = subject.name
            cl+=1
        doc.add_page_break()

    document_io = io.BytesIO()
    doc.save(document_io)
    document_io.seek(0)
    response = FileResponse(document_io, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{request.user.schooladministrator.current_school}_{exam}_assessment_sheets.docx"'
    return response


def exam_papers(request, exam):
    exam = request.user.schooladministrator.current_school.Examinations.get(id=exam)
    context = {
        'exam': exam
    }
    return render(request, 'exam_papers.html', context)


def schedule_paper(request, exam):
    form = PaperForm()
    exam = request.user.schooladministrator.current_school.Examinations.get(id=exam)
    form.fields['Subject'].queryset = request.user.schooladministrator.current_school.Subjects.all()
    form.fields['Class'].queryset = request.user.schooladministrator.current_school.classes.all()
    form.fields['Examiner'].queryset = request.user.schooladministrator.current_school.Teachers.all()
    context = {
        'form': form
    }
    if request.method == 'POST':
        form = PaperForm(request.POST)
        if form.is_valid():
            form = PaperForm(request.POST)
            paper = form.save(commit=False)
            paper.Examination = exam
            paper.save()
            exam.Papers.add(paper)
            return redirect('exam-papers', exam=exam.id)
    else:
        return render(request, 'schedule_paper.html', context)