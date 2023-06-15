from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import Report
from .forms import BatchReportsForm, PublishBatchReportsForm
import random
from django.urls import reverse
from Students.models import Student
from Examinations.models import Examination
import docx
from django.core.mail import send_mail
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Image, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import zipfile
from django.http import FileResponse, HttpResponse


def reports(request):
    reports = Report.objects.all()
    context = {
        'reports': reports
    }
    return render(request, 'reports.html', context)


def student_reports(request, student):
    student = Student.objects.get(id=student)
    context = {
        'student': student
    }
    return render(request, 'student_reports.html', context)


def edit_report(request, report):
    report = Report.objects.get(pk=report)
    school = request.user.schooladministrator.current_school
    context = {
        'report': report    
    }
    if request.method == 'POST':
        for score in report.Scores.all():
            sc = request.POST[str(score.pk)]
            grade = request.POST[f'{score.pk}grade']
            grade = school.Grades.get(pk=grade)
            score.Grade = grade
            score.Score = sc
            score.save()
        report.Total_Score = sum([sc.Score for sc in report.Scores.all()])
        report.Aggregate = sum([sc.Grade.Value for sc in report.Scores.all()])
        report.save()
        return redirect(request.session['next'])
    else:
        request.session['next'] = request.META.get('HTTP_REFERER')
        return render(request, 'edit_report.html', context)


def view_report(request):
    return render(request, 'report.html')


@login_required
def delete_report(request, pk):
    report = get_object_or_404(Report, pk=pk)
    if request.method == 'POST':
        report.delete()
        messages.success(request, 'Report deleted successfully.')
        return redirect('report_list')
    context = {'report': report}
    return render(request, 'delete_report.html', context)


def create_report(request, student, examination):
    student = Student.objects.get(pk=student)
    exam = Examination.objects.get(pk=examination)
    report = student.Reports.create(Student=student, Examination=exam)
    for subject in student.Subjects.all():
        report.Scores.create(Score=0, Report=report, Subject=subject)
    student.school.Reports.add(report)
    return redirect('edit-report', report=report.id)


def download_report(request, report):
    report = Report.objects.get(pk=report)
    doc = docx.Document()
    for clas in request.user.schooladministrator.current_school.classes.all():
        doc.add_heading(f'{report.student} {report.student.Examination} Reports', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        table = doc.add_table(rows=report.Student.Subjects.count()+4, cols=4)
        table.cell(0, 0).text = 'Subject'
        table.cell(0, 1).text = 'Marks'
        cl = 1
        for score in report.Scores.all():
            table.cell(0, cl).text = score.Subject.name
            cl+=1
        doc.add_page_break()

    document_io = io.BytesIO()
    doc.save(document_io)
    document_io.seek(0)
    response = FileResponse(document_io, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{request.user.schooladministrator.current_school}_{exam}_assessment_sheets.docx"'
    return response

def publish_batch(request):
    form = PublishBatchReportsForm()
    form.fields['Exam'].queryset = request.user.schooladministrator.current_school.Examinations.all()
    form.fields['Classes'].queryset = request.user.schooladministrator.current_school.classes.all()
    context = {
        'form': form
    }
    if request.method == 'POST':
        form = PublishBatchReportsForm(request.POST)
        if form.is_valid():
            data = form.cleaned_data
            reports = request.user.schooladministrator.current_school.Reports.filter(Student__Class__in=data['Classes'], Examination=data['Exam'])
            parents = request.user.schooladministrator.current_school.Parents.filter(relationships__Child__Class__in=data['Classes']).distinct()
            emails = [p.user.email for p in parents]
            message = f'''
            {request.user.schooladministrator.current_school} {data["Exam"]} Reports have been published please login as a parent and checkout the reports section.
            '''
            send_mail(
                f'{data["Exam"]} Reports Published',
                message,
                'Edumetrics <info@edu-metrics.com>',
                emails
            )
            messages.success(request, 'Reports have been published successfully.')
            reports.update(Published=True)
            return redirect('reports')
    else:
        return render(request, 'publish_batch.html', context)


def child_reports(request):
    return render(request, 'child_reports.html')

def export_reports(request):
    file_paths = []
    for report in request.user.schooladministrator.current_school.Reports.all():
        doc = SimpleDocTemplate(f"{report.id}.pdf", pagesize=A4)
        file_paths.append(f"{report.id}.pdf")

        # Create a list to hold the flowables (elements) of the document
        elements = []

        # Define the styles for the document
        styles = getSampleStyleSheet()
        title_style = styles["Title"]
        table_style = TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), "red"),
                ("TEXTCOLOR", (0, 0), (-1, 0), "white"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 14),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                ("TEXTCOLOR", (0, 1), (-1, -1), "black"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 1), (-1, -1), 12),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("GRID", (0, 0), (-1, -1), 1, "black"),  # Add border to all cells
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),  # Center text vertically in cells
            ]
        )
        try:
            image_path1 = report.Student.photo.path
            image_path2 = report.Student.photo.path
            image1 = Image(image_path1, width=200, height=200)
            image2 = Image(image_path1, width=200, height=200)
            elements.append(image1)
        except:
            pass

        # Create a title for the document
        title = Paragraph(str(request.user.schooladministrator.current_school), title_style)
        elements.append(title)

        # Create a table to display the results
        data = [["Subject", "Score", "Grade", "Comment"]]
        for result in report.Results.all():
            data.append([result.Paper.Subject.name, result.Score, '', ''])

        table = Table(data)
        table.setStyle(table_style)
        elements.append(table)

        # Build the PDF document
        doc.build(elements)
    zip_name = 'reports.zip'
    with zipfile.ZipFile(zip_name, 'w') as zipf:
        for file_path in file_paths:
            zipf.write(file_path)
    with open(zip_name, 'rb') as zip_file:
        response = HttpResponse(zip_file.read(), content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="{zip_name}"'
        return response